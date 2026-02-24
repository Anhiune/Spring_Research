"""
Converts research_progress_report.md to a Word .docx
- Removes all em-dashes (replaces contextually with colon, comma, or parenthesis)
- Rewrites third-person / second-person language to first person
- Preserves tables, bullet lists, bold, code blocks, headers
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
BASE  = Path(__file__).parent
IN_MD = BASE / "research_progress_report.md"
OUT   = BASE / "research_progress_report.docx"

# ── Step 1: Load and apply all text transformations ────────────────────────

text = IN_MD.read_text(encoding="utf-8")

# ---- 1a. First-person substitutions (order matters — specific before general) ----
fp = [
    # Title / meta
    ("What Has Been Accomplished",              "What I Have Accomplished"),
    # Section intros
    ("A full automated Bluesky scraper was built from scratch",
     "I built a full automated Bluesky scraper from scratch"),
    ("Keywords are organized into five thematic groups",
     "I organized the keywords into five thematic groups"),
    ("The background scrape is still running.",
     "My background scrape is still running."),
    ("The background scrape is still running",
     "My background scrape is still running"),
    ("The scraper is continuing in the background",
     "The scraper continues in the background"),
    # Cleaning section
    ("Cleaning must happen before any feature engineering or modeling. Work through these\nsteps in order.",
     "I need to complete cleaning before any feature engineering or modeling, and I will work through these steps in order."),
    ("Cleaning must happen before any feature engineering or modeling. Work through these steps in order.",
     "I need to complete cleaning before any feature engineering or modeling, and I will work through these steps in order."),
    ("Even though the scraper deduplicates within each month, the same post could\ntheoretically appear across two monthly files at a month boundary.",
     "Even though the scraper deduplicates within each month, the same post could theoretically appear across two monthly files at a month boundary."),
    ("Export-control discourse is predominantly English. Non-English posts add noise.",
     "Export-control discourse is predominantly English, so non-English posts add noise to my data."),
    ("Not all posts mentioning \"Intel\" or \"AMD\" are about export controls.\nYou need a relevance filter to separate:",
     "Not all posts mentioning \"Intel\" or \"AMD\" are about export controls.\nI need a relevance filter to separate:"),
    ("Not all posts mentioning \"Intel\" or \"AMD\" are about export controls.",
     "Not all posts mentioning \"Intel\" or \"AMD\" are about export controls."),
    ("You need a relevance filter to separate:",
     "I need a relevance filter to separate:"),
    ("Bluesky posts are timestamped in UTC. You need to map each post to the correct",
     "Bluesky posts are timestamped in UTC. I need to map each post to the correct"),
    # Section 4
    ("After cleaning, build daily sentiment indices from the post text.",
     "After cleaning, I will build daily sentiment indices from the post text."),
    ("Since your existing Tesla model uses the **NRC emotion lexicon** (Bollen method),\nuse the same approach for continuity",
     "Since my existing Tesla model uses the **NRC emotion lexicon** (Bollen method),\nI will use the same approach for continuity"),
    ("Since your existing Tesla model uses",
     "Since my existing Tesla model uses"),
    ("use the same approach for continuity",
     "I will use the same approach for continuity"),
    ("**Recommendation:** Run both. Use FinBERT as primary (it understands financial jargon\nlike \"license denied\", \"entity list\") and NRC as a robustness check.",
     "**My recommendation:** I will run both, using FinBERT as the primary scorer (it understands financial jargon like \"license denied\" and \"entity list\") and NRC as a robustness check."),
    ("Separate from the general sentiment, build a **daily policy-risk index** using\nonly Tier 1 (policy-context) posts. This is the key explanatory variable for H1 and H2.",
     "Separately, I will build a **daily policy-risk index** using only Tier 1 (policy-context) posts. This is the key explanatory variable for H1 and H2."),
    # Section 5
    ("Manually score each firm (0",
     "I will manually score each firm (0"),
    ("Build your policy-event calendar (BIS rule dates, Entity List additions, major\n   Congressional actions).",
     "I will build my policy-event calendar (BIS rule dates, Entity List additions, and major Congressional actions)."),
    ("Build your policy-event calendar",
     "I will build my policy-event calendar"),
    ("For each event, compute Cumulative Abnormal Returns (CARs) in windows",
     "For each event, I will compute Cumulative Abnormal Returns (CARs) in windows"),
    ("Test whether pre-event policy-risk sentiment (averaged over days [-10, -1])\n   predicts the size of CARs. Run:",
     "I will then test whether pre-event policy-risk sentiment (averaged over days [-10, -1]) predicts the size of CARs by running:"),
    ("Test whether pre-event policy-risk sentiment",
     "I will test whether pre-event policy-risk sentiment"),
    ("Extend your existing **ARIMAX model** from the Submition folder to the new context:",
     "I will extend my existing **ARIMAX model** from the Submition folder to the new context:"),
    ("Where `realized_vol_{t+5}` = standard deviation of log returns over the next 5 trading days.",
     "where `realized_vol_{t+5}` is the standard deviation of log returns over the next 5 trading days."),
    ("The interaction term `exposure_score × L1_policy_risk_sentiment` directly tests H3\n(heterogeneous effects).",
     "The interaction term `exposure_score x L1_policy_risk_sentiment` directly tests H3 (heterogeneous effects)."),
    # Section 6
    ("The existing model in `Submition/` uses Tesla tweets as its sentiment source.\nHere are concrete extensions ranked by priority:",
     "My existing model in `Submition/` uses Tesla tweets as its sentiment source.\nBelow are concrete extensions I have ranked by priority:"),
    ("The current `sentiment_index` comes from Tesla tweets via the Twitter v2 API.\nYou can now augment it with Bluesky data (for 2023 onward), using the same NRC\nlexicon method already implemented in `Practice with Bollen method_sentiment_analysis_step3.py`.\nThis adds a **second, independent social media signal** and allows you to test\nwhether cross-platform sentiment agreement is more predictive than either alone.",
     "My current `sentiment_index` comes from Tesla tweets via the Twitter v2 API.\nI can now augment it with Bluesky data (for 2023 onward), using the same NRC\nlexicon method already implemented in `Practice with Bollen method_sentiment_analysis_step3.py`.\nThis adds a **second, independent social media signal** and allows me to test\nwhether cross-platform sentiment agreement is more predictive than either alone."),
    ("The current ARIMAX has `L1_sentiment` and `L2_sentiment` but no explicit policy-risk signal.\nAdd a `policy_risk_index` (built from Tier 1 posts) as a new regressor:",
     "My current ARIMAX has `L1_sentiment` and `L2_sentiment` but no explicit policy-risk signal.\nI will add a `policy_risk_index` (built from Tier 1 posts) as a new regressor:"),
    ("Compare AIC/BIC between the old and new models. If `L1_policy_risk` is significant,\nthat directly answers H2.",
     "I will compare AIC/BIC between the old and new models. If `L1_policy_risk` is significant, this directly answers H2."),
    ("Your current model predicts daily **returns**. Switch (or add) **realized volatility**\nas the dependent variable. This is more consistent with H2 and better-suited to\nARIMAX (volatility is roughly stationary without differencing, and the ARCH test you\nalready run on residuals motivates this).",
     "My current model predicts daily **returns**. I will switch to (or add) **realized volatility** as the dependent variable, which is more consistent with H2 and better suited to ARIMAX since volatility is roughly stationary without differencing, and the ARCH test I already run on residuals motivates this choice."),
    ("Your existing residuals already show ARCH effects (you run `ArchTest()`). The natural\nnext step is an **ARIMA-GARCH** or **ARIMAX-GARCH** model to model both the conditional\nmean and conditional variance:",
     "My existing residuals already show ARCH effects (from the `ArchTest()` I run). The natural next step is an **ARIMA-GARCH** or **ARIMAX-GARCH** model that models both the conditional mean and conditional variance:"),
    ("Currently the model is estimated in-sample. Add a rolling or expanding window\nout-of-sample evaluation:",
     "Currently my model is estimated in-sample. I will add a rolling out-of-sample evaluation:"),
    ("You noted the possibility of adding Google Trends. This is a useful proxy for\nretail/public attention:",
     "I noted the possibility of adding Google Trends, which is a useful proxy for retail and public attention:"),
    ("Google Trends provides a weekly or daily normalized Search Volume Index (SVI) that\nyou can merge by date and include alongside Bluesky sentiment.",
     "Google Trends provides a weekly or daily normalized Search Volume Index (SVI) that I can merge by date and include alongside Bluesky sentiment."),
    # Checklist
    ("Let batch scraper finish", "Let my batch scraper finish"),
    ("Run cleaning steps", "Run my cleaning steps"),
    # Footer
    ("Document auto-generated from workspace analysis",
     "Document prepared from workspace analysis"),
    # Generic second-person cleanup
    ("your research question",   "my research question"),
    ("your existing",            "my existing"),
    ("your policy-event",        "my policy-event"),
    ("# In R — consistent with your existing ARIMAX code style",
     "# In R, consistent with my existing ARIMAX code style"),
]

for old, new in fp:
    text = text.replace(old, new)

# ---- 1b. Remove em-dashes contextually ----
# Section headers: "### X.Y Title — Subtitle"  ->  "### X.Y Title: Subtitle"
text = re.sub(r'(#{1,4} [^\n]+?) — ([^\n]+)', r'\1: \2', text)

# Bullet list items: "- **Label** — description"  ->  "- **Label**: description"
text = re.sub(r'^(- \*\*[^*]+\*\*) — ', r'\1: ', text, flags=re.MULTILINE)

# Table cells that use " — " as a separator within the cell
text = re.sub(r' — ', ', ', text)

# Any remaining stray em-dashes not in code blocks
lines = text.split('\n')
cleaned = []
in_code = False
for line in lines:
    if line.startswith('```'):
        in_code = not in_code
    if not in_code:
        line = line.replace('—', '-')
    cleaned.append(line)
text = '\n'.join(cleaned)

# ── Step 2: Parse processed markdown into a structured list of blocks ──────

def parse_inline(text_str):
    """Return list of (string, bold, mono) tuples from inline markdown."""
    segments = []
    # Split on **bold** and `mono` patterns
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text_str)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], True, False))
        elif part.startswith('`') and part.endswith('`'):
            segments.append((part[1:-1], False, True))
        else:
            segments.append((part, False, False))
    return segments


def parse_table(lines_block):
    """Parse markdown table lines into list of rows (list of cells)."""
    rows = []
    for line in lines_block:
        line = line.strip()
        if not line or set(line.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
            continue  # skip separator lines
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


# ── Step 3: Build Word document ────────────────────────────────────────────

doc = Document()

# Document-wide style defaults
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title page info
lines_all = text.split('\n')

i = 0
while i < len(lines_all):
    line = lines_all[i]

    # ---- Horizontal rule ----
    if line.strip() == '---':
        # Add a thin paragraph as a visual break
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        i += 1
        continue

    # ---- Headings ----
    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    if line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
        i += 1
        continue

    # ---- Code block ----
    if line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines_all) and not lines_all[i].startswith('```'):
            code_lines.append(lines_all[i])
            i += 1
        i += 1  # skip closing ```
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph(style='No Spacing')
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        continue

    # ---- Blockquote (>) ----
    if line.startswith('> '):
        content = line[2:].strip()
        p = doc.add_paragraph(style='No Spacing')
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        for seg, bold, mono in parse_inline(content):
            run = p.add_run(seg)
            run.bold = bold
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        i += 1
        continue

    # ---- Table ----
    if line.startswith('|'):
        table_lines = []
        while i < len(lines_all) and lines_all[i].startswith('|'):
            table_lines.append(lines_all[i])
            i += 1
        rows = parse_table(table_lines)
        if not rows:
            continue
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = 'Table Grid'
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                for seg, bold, mono in parse_inline(cell_text):
                    run = p.add_run(seg)
                    run.bold = bold or (r_idx == 0)  # header row always bold
                    if mono:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
        doc.add_paragraph()  # spacer after table
        continue

    # ---- Bullet list item ----
    if line.startswith('- ') or line.startswith('  - '):
        indent = 1 if line.startswith('  - ') else 0
        content = line.lstrip('- ').strip()
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        for seg, bold, mono in parse_inline(content):
            run = p.add_run(seg)
            run.bold = bold
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        i += 1
        continue

    # ---- Numbered list item ----
    m = re.match(r'^(\d+)\. (.+)', line)
    if m:
        content = m.group(2)
        p = doc.add_paragraph(style='List Number')
        for seg, bold, mono in parse_inline(content):
            run = p.add_run(seg)
            run.bold = bold
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        i += 1
        continue

    # ---- Bold metadata lines (**Key:** Value) ----
    if line.startswith('**') and not line.startswith('**Option') and ':**' in line:
        p = doc.add_paragraph(style='No Spacing')
        for seg, bold, mono in parse_inline(line):
            run = p.add_run(seg)
            run.bold = bold
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        i += 1
        continue

    # ---- Empty line ----
    if line.strip() == '':
        i += 1
        continue

    # ---- Regular paragraph ----
    p = doc.add_paragraph()
    for seg, bold, mono in parse_inline(line):
        run = p.add_run(seg)
        run.bold = bold
        if mono:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    i += 1

# ── Step 4: Save ────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
